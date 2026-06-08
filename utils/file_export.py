"""
CV Tailor — File Export Utilities
==================================
Generate PDF and DOCX files from structured resume data.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from config import OUTPUTS_DIR, get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Resume Data Model (dict-based for flexibility)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

"""
Expected resume_data dict structure:
{
    "name": str,
    "email": str,
    "phone": str,
    "location": str,
    "linkedin": str,          # optional
    "website": str,            # optional
    "summary": str,
    "skills": list[str],
    "experience": [
        {
            "title": str,
            "company": str,
            "location": str,
            "start_date": str,
            "end_date": str,
            "bullets": list[str],
        }
    ],
    "education": [
        {
            "degree": str,
            "institution": str,
            "location": str,
            "graduation_date": str,
            "gpa": str,            # optional
        }
    ],
    "certifications": list[str],   # optional
    "projects": [                   # optional
        {
            "name": str,
            "description": str,
            "technologies": str,
        }
    ],
    "languages": list[str],        # optional
}
"""


def generate_filename(name: str, ext: str) -> str:
    """Generate a timestamped filename.

    Args:
        name: Person's name for the filename.
        ext: File extension (pdf, docx).

    Returns:
        Filename string.
    """
    safe_name = name.replace(" ", "_").lower()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{safe_name}_resume_{timestamp}.{ext}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  DOCX Export
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def export_docx(resume_data: dict[str, Any], output_path: str | None = None) -> Path:
    """Export resume data to a DOCX file.

    Args:
        resume_data: Structured resume dictionary.
        output_path: Optional custom output path.

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Name / Header ─────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data.get("name", "Your Name"))
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Contact line
    contact_parts: list[str] = []
    if resume_data.get("email"):
        contact_parts.append(resume_data["email"])
    if resume_data.get("phone"):
        contact_parts.append(resume_data["phone"])
    if resume_data.get("location"):
        contact_parts.append(resume_data["location"])
    if resume_data.get("linkedin"):
        contact_parts.append(resume_data["linkedin"])
    if resume_data.get("website"):
        contact_parts.append(resume_data["website"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run("  |  ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_section_heading(title: str) -> None:
        """Add a styled section heading with a bottom border."""
        para = doc.add_paragraph()
        para.space_before = Pt(12)
        para.space_after = Pt(4)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # Add bottom border via XML
        from docx.oxml.ns import qn
        from lxml import etree

        pPr = para._p.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A1A2E")

    # ── Summary ───────────────────────────────────────────────
    if resume_data.get("summary"):
        _add_section_heading("Professional Summary")
        summary_para = doc.add_paragraph(resume_data["summary"])
        summary_para.paragraph_format.space_after = Pt(2)

    # ── Skills ────────────────────────────────────────────────
    if resume_data.get("skills"):
        _add_section_heading("Skills")
        skills_text = "  •  ".join(resume_data["skills"])
        skills_para = doc.add_paragraph(skills_text)
        skills_para.paragraph_format.space_after = Pt(2)

    # ── Experience ────────────────────────────────────────────
    if resume_data.get("experience"):
        _add_section_heading("Experience")
        for exp in resume_data["experience"]:
            # Title + Company
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(6)
            title_para.paragraph_format.space_after = Pt(1)
            title_run = title_para.add_run(exp.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            title_para.add_run(f"  –  {exp.get('company', '')}")

            # Location + Dates
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(2)
            date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            loc = exp.get("location", "")
            meta = f"{loc}  |  {date_str}" if loc else date_str
            date_run = date_para.add_run(meta)
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            date_run.italic = True

            # Bullets
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = bullet
                bp.paragraph_format.space_after = Pt(1)

    # ── Education ─────────────────────────────────────────────
    if resume_data.get("education"):
        _add_section_heading("Education")
        for edu in resume_data["education"]:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(4)
            edu_para.paragraph_format.space_after = Pt(1)
            deg_run = edu_para.add_run(edu.get("degree", ""))
            deg_run.bold = True
            edu_para.add_run(f"  –  {edu.get('institution', '')}")

            meta_parts: list[str] = []
            if edu.get("location"):
                meta_parts.append(edu["location"])
            if edu.get("graduation_date"):
                meta_parts.append(edu["graduation_date"])
            if edu.get("gpa"):
                meta_parts.append(f"GPA: {edu['gpa']}")
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_para.paragraph_format.space_after = Pt(2)
                meta_run = meta_para.add_run("  |  ".join(meta_parts))
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                meta_run.italic = True

    # ── Certifications ────────────────────────────────────────
    if resume_data.get("certifications"):
        _add_section_heading("Certifications")
        for cert in resume_data["certifications"]:
            doc.add_paragraph(f"• {cert}", style="List Bullet")

    # ── Projects ──────────────────────────────────────────────
    if resume_data.get("projects"):
        _add_section_heading("Projects")
        for proj in resume_data["projects"]:
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(4)
            proj_run = proj_para.add_run(proj.get("name", ""))
            proj_run.bold = True
            proj_tech = proj.get("technologies", "").strip()
            if proj_tech:
                tech_run = proj_para.add_run(f"  ({proj_tech})")
                tech_run.italic = True
                tech_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if proj.get("description"):
                doc.add_paragraph(proj["description"])

    # ── Languages ─────────────────────────────────────────────
    if resume_data.get("languages"):
        _add_section_heading("Languages")
        doc.add_paragraph("  •  ".join(resume_data["languages"]))

    # ── Research Papers ───────────────────────────────────────
    if resume_data.get("research_papers"):
        _add_section_heading("Research & Publications")
        for paper in resume_data["research_papers"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run(paper.get('title', ''))
            r.bold = True
            authors = paper.get('authors', '').strip()
            if authors:
                p.add_run(f" — {authors}")
            meta = []
            if paper.get('venue'): meta.append(paper['venue'])
            if paper.get('year'): meta.append(paper['year'])
            if paper.get('url'): meta.append(paper['url'])
            if meta:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mr = mp.add_run(' | '.join(meta))
                mr.italic = True
                mr.font.size = Pt(9)
                mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Volunteer Work ───────────────────────────────────────
    if resume_data.get("volunteer_work"):
        _add_section_heading("Volunteer Experience")
        for vol in resume_data["volunteer_work"]:
            vp = doc.add_paragraph()
            vp.paragraph_format.space_before = Pt(4)
            vr = vp.add_run(vol.get('role', ''))
            vr.bold = True
            org = vol.get('organization', '').strip()
            if org:
                vp.add_run(f" — {org}")
            dates = []
            if vol.get('start_date'): dates.append(vol['start_date'])
            if vol.get('end_date'): dates.append(vol['end_date'])
            if dates:
                dp = doc.add_paragraph()
                dp.paragraph_format.space_before = Pt(0)
                dr = dp.add_run(' — '.join(dates))
                dr.italic = True
                dr.font.size = Pt(9)
                dr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if vol.get('description'):
                doc.add_paragraph(vol['description'])

    # ── Achievements ─────────────────────────────────────────
    if resume_data.get("achievements"):
        _add_section_heading("Achievements & Awards")
        for ach in resume_data["achievements"]:
            doc.add_paragraph(f"• {ach}", style="List Bullet")

    # ── Save ──────────────────────────────────────────────────
    if output_path is None:
        fname = generate_filename(resume_data.get("name", "resume"), "docx")
        output_path = str(OUTPUTS_DIR / fname)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("DOCX exported: %s", output_path)
    return Path(output_path)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  PDF Export
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def export_pdf(resume_data: dict[str, Any], output_path: str | None = None) -> Path:
    """Export resume data to a PDF file using ReportLab.

    Args:
        resume_data: Structured resume dictionary.
        output_path: Optional custom output path.

    Returns:
        Path to the generated PDF file.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        HRFlowable,
    )

    if output_path is None:
        fname = generate_filename(resume_data.get("name", "resume"), "pdf")
        output_path = str(OUTPUTS_DIR / fname)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name="ResumeName",
        parent=styles["Title"],
        fontSize=22,
        leading=26,
        textColor=HexColor("#1A1A2E"),
        alignment=1,  # center
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="ContactLine",
        parent=styles["Normal"],
        fontSize=9,
        textColor=HexColor("#666666"),
        alignment=1,
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name="SectionHeading",
        parent=styles["Heading2"],
        fontSize=11,
        textColor=HexColor("#1A1A2E"),
        spaceBefore=14,
        spaceAfter=4,
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="ResumeBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        textColor=HexColor("#333333"),
        spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name="ResumeMeta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=HexColor("#666666"),
        spaceAfter=3,
        fontName="Helvetica-Oblique",
    ))
    styles.add(ParagraphStyle(
        name="BulletPoint",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        textColor=HexColor("#333333"),
        leftIndent=18,
        spaceAfter=2,
        bulletIndent=6,
    ))

    elements: list[Any] = []
    divider = HRFlowable(width="100%", thickness=1, color=HexColor("#1A1A2E"), spaceAfter=6)

    # ── Name ──
    elements.append(Paragraph(resume_data.get("name", "Your Name"), styles["ResumeName"]))

    # ── Contact ──
    contact_parts: list[str] = []
    for key in ("email", "phone", "location", "linkedin", "website"):
        if resume_data.get(key):
            contact_parts.append(resume_data[key])
    if contact_parts:
        elements.append(Paragraph("  |  ".join(contact_parts), styles["ContactLine"]))

    # ── Summary ──
    if resume_data.get("summary"):
        elements.append(Paragraph("PROFESSIONAL SUMMARY", styles["SectionHeading"]))
        elements.append(divider)
        elements.append(Paragraph(resume_data["summary"], styles["ResumeBody"]))

    # ── Skills ──
    if resume_data.get("skills"):
        elements.append(Paragraph("SKILLS", styles["SectionHeading"]))
        elements.append(divider)
        skills_text = "  •  ".join(resume_data["skills"])
        elements.append(Paragraph(skills_text, styles["ResumeBody"]))

    # ── Experience ──
    if resume_data.get("experience"):
        elements.append(Paragraph("EXPERIENCE", styles["SectionHeading"]))
        elements.append(divider)
        for exp in resume_data["experience"]:
            title_text = f"<b>{exp.get('title', '')}</b>  –  {exp.get('company', '')}"
            elements.append(Paragraph(title_text, styles["ResumeBody"]))
            date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            loc = exp.get("location", "")
            meta = f"{loc}  |  {date_str}" if loc else date_str
            elements.append(Paragraph(meta, styles["ResumeMeta"]))
            for bullet in exp.get("bullets", []):
                elements.append(Paragraph(f"•  {bullet}", styles["BulletPoint"]))
            elements.append(Spacer(1, 4))

    # ── Education ──
    if resume_data.get("education"):
        elements.append(Paragraph("EDUCATION", styles["SectionHeading"]))
        elements.append(divider)
        for edu in resume_data["education"]:
            deg_text = f"<b>{edu.get('degree', '')}</b>  –  {edu.get('institution', '')}"
            elements.append(Paragraph(deg_text, styles["ResumeBody"]))
            meta_parts: list[str] = []
            if edu.get("location"):
                meta_parts.append(edu["location"])
            if edu.get("graduation_date"):
                meta_parts.append(edu["graduation_date"])
            if edu.get("gpa"):
                meta_parts.append(f"GPA: {edu['gpa']}")
            if meta_parts:
                elements.append(Paragraph("  |  ".join(meta_parts), styles["ResumeMeta"]))

    # ── Certifications ──
    if resume_data.get("certifications"):
        elements.append(Paragraph("CERTIFICATIONS", styles["SectionHeading"]))
        elements.append(divider)
        for cert in resume_data["certifications"]:
            elements.append(Paragraph(f"•  {cert}", styles["BulletPoint"]))

    # ── Projects ──
    if resume_data.get("projects"):
        elements.append(Paragraph("PROJECTS", styles["SectionHeading"]))
        elements.append(divider)
        for proj in resume_data["projects"]:
            proj_name = proj.get('name', '')
            proj_tech = proj.get('technologies', '').strip()
            proj_text = f"<b>{proj_name}</b>"
            if proj_tech:
                proj_text += f"  <i>({proj_tech})</i>"
            elements.append(Paragraph(proj_text, styles["ResumeBody"]))
            if proj.get("description"):
                elements.append(Paragraph(proj["description"], styles["ResumeBody"]))

    # ── Languages ──
    if resume_data.get("languages"):
        elements.append(Paragraph("LANGUAGES", styles["SectionHeading"]))
        elements.append(divider)
        elements.append(Paragraph("  •  ".join(resume_data["languages"]), styles["ResumeBody"]))

    # ── Research Papers ──
    if resume_data.get("research_papers"):
        elements.append(Paragraph("RESEARCH & PUBLICATIONS", styles["SectionHeading"]))
        elements.append(divider)
        for paper in resume_data["research_papers"]:
            title = paper.get('title', '')
            authors = paper.get('authors', '')
            venue = paper.get('venue', '')
            year = paper.get('year', '')
            url = paper.get('url', '')
            txt = f"<b>{title}</b>"
            if authors: txt += f" — {authors}"
            meta = []
            if venue: meta.append(f"<i>{venue}</i>")
            if year: meta.append(year)
            if url: meta.append(url)
            if meta: txt += f"<br/>{'  |  '.join(meta)}"
            elements.append(Paragraph(txt, styles["ResumeBody"]))

    # ── Volunteer Work ──
    if resume_data.get("volunteer_work"):
        elements.append(Paragraph("VOLUNTEER EXPERIENCE", styles["SectionHeading"]))
        elements.append(divider)
        for vol in resume_data["volunteer_work"]:
            role = vol.get('role', '')
            org = vol.get('organization', '')
            txt = f"<b>{role}</b>"
            if org: txt += f" — {org}"
            elements.append(Paragraph(txt, styles["ResumeBody"]))
            dates = []
            if vol.get('start_date'): dates.append(vol['start_date'])
            if vol.get('end_date'): dates.append(vol['end_date'])
            if dates:
                elements.append(Paragraph(' — '.join(dates), styles["ResumeMeta"]))
            if vol.get('description'):
                elements.append(Paragraph(vol['description'], styles["ResumeBody"]))

    # ── Achievements ──
    if resume_data.get("achievements"):
        elements.append(Paragraph("ACHIEVEMENTS & AWARDS", styles["SectionHeading"]))
        elements.append(divider)
        for ach in resume_data["achievements"]:
            elements.append(Paragraph(f"•  {ach}", styles["BulletPoint"]))

    doc.build(elements)
    logger.info("PDF exported: %s", output_path)
    return Path(output_path)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Plain-text export (for tailoring / processing)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def resume_data_to_text(resume_data: dict[str, Any]) -> str:
    """Convert structured resume data to plain text.

    Args:
        resume_data: Structured resume dictionary.

    Returns:
        Plain-text resume string.
    """
    lines: list[str] = []

    lines.append(resume_data.get("name", ""))
    contact_parts = []
    for key in ("email", "phone", "location", "linkedin", "website"):
        if resume_data.get(key):
            contact_parts.append(resume_data[key])
    if contact_parts:
        lines.append(" | ".join(contact_parts))
    lines.append("")

    if resume_data.get("summary"):
        lines.append("Summary")
        lines.append(resume_data["summary"])
        lines.append("")

    if resume_data.get("skills"):
        lines.append("Skills")
        lines.append(", ".join(resume_data["skills"]))
        lines.append("")

    if resume_data.get("experience"):
        lines.append("Experience")
        for exp in resume_data["experience"]:
            title = exp.get('title', '').strip()
            company = exp.get('company', '').strip()
            if title and company:
                lines.append(f"{title} - {company}")
            elif title:
                lines.append(title)
            elif company:
                lines.append(company)
            date_parts = []
            loc = exp.get("location", "").strip()
            if loc:
                date_parts.append(loc)
            start = exp.get('start_date', '').strip()
            end = exp.get('end_date', 'Present').strip()
            if start:
                date_parts.append(f"{start} - {end}")
            if date_parts:
                lines.append("  |  ".join(date_parts))
            for bullet in exp.get("bullets", []):
                lines.append(f"  - {bullet}")
            lines.append("")

    if resume_data.get("education"):
        lines.append("Education")
        for edu in resume_data["education"]:
            degree = edu.get('degree', '').strip()
            institution = edu.get('institution', '').strip()
            if degree and institution:
                lines.append(f"{degree} - {institution}")
            elif degree:
                lines.append(degree)
            elif institution:
                lines.append(institution)
            meta_parts = []
            if edu.get("location"):
                meta_parts.append(edu["location"])
            if edu.get("graduation_date"):
                meta_parts.append(edu["graduation_date"])
            if edu.get("gpa"):
                meta_parts.append(f"GPA: {edu['gpa']}")
            if meta_parts:
                lines.append("  |  ".join(meta_parts))
            lines.append("")

    if resume_data.get("certifications"):
        lines.append("Certifications")
        for cert in resume_data["certifications"]:
            lines.append(f"  - {cert}")
        lines.append("")

    if resume_data.get("projects"):
        lines.append("Projects")
        for proj in resume_data["projects"]:
            proj_name = proj.get('name', '')
            proj_tech = proj.get('technologies', '').strip()
            if proj_tech:
                lines.append(f"{proj_name} ({proj_tech})")
            else:
                lines.append(proj_name)
            if proj.get("description"):
                lines.append(f"  {proj['description']}")
        lines.append("")

    if resume_data.get("languages"):
        lines.append("Languages")
        lines.append(", ".join(resume_data["languages"]))
        lines.append("")

    if resume_data.get("research_papers"):
        lines.append("Research & Publications")
        for paper in resume_data["research_papers"]:
            title = paper.get('title', '')
            authors = paper.get('authors', '').strip()
            line = title
            if authors:
                line += f" - {authors}"
            lines.append(line)
            meta = []
            if paper.get('venue'): meta.append(paper['venue'])
            if paper.get('year'): meta.append(paper['year'])
            if paper.get('url'): meta.append(paper['url'])
            if meta:
                lines.append("  " + "  |  ".join(meta))
        lines.append("")

    if resume_data.get("volunteer_work"):
        lines.append("Volunteer Experience")
        for vol in resume_data["volunteer_work"]:
            role = vol.get('role', '')
            org = vol.get('organization', '').strip()
            if role and org:
                lines.append(f"{role} - {org}")
            elif role:
                lines.append(role)
            dates = []
            if vol.get('start_date'): dates.append(vol['start_date'])
            if vol.get('end_date'): dates.append(vol['end_date'])
            if dates:
                lines.append(" - ".join(dates))
            if vol.get('description'):
                lines.append(f"  {vol['description']}")
        lines.append("")

    if resume_data.get("achievements"):
        lines.append("Achievements & Awards")
        for ach in resume_data["achievements"]:
            lines.append(f"  - {ach}")
        lines.append("")

    return "\n".join(lines).strip()


def export_text(resume_data: dict[str, Any], output_path: str | None = None) -> Path:
    """Export resume data to a plain text file.

    Args:
        resume_data: Structured resume dictionary.
        output_path: Optional custom output path.

    Returns:
        Path to the generated text file.
    """
    text = resume_data_to_text(resume_data)

    if output_path is None:
        fname = generate_filename(resume_data.get("name", "resume"), "txt")
        output_path = str(OUTPUTS_DIR / fname)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(text, encoding="utf-8")
    logger.info("TXT exported: %s", output_path)
    return Path(output_path)
