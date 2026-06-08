"""
CV Tailor — Resume Templates Module
======================================
Provides multiple ATS-friendly resume templates for PDF and DOCX export.
Templates: minimal, professional, modern.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Literal

from config import TEMPLATES_DIR, OUTPUTS_DIR

logger = logging.getLogger(__name__)

TemplateName = Literal["minimal", "professional", "modern"]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Template Configuration
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TEMPLATE_CONFIGS: dict[str, dict[str, Any]] = {
    "minimal": {
        "name": "Minimal",
        "description": "Clean, single-column layout with minimal styling.",
        "font": "Helvetica",
        "heading_font": "Helvetica-Bold",
        "font_size": 10,
        "heading_size": 11,
        "name_size": 18,
        "section_spacing": 10,
        "bullet_spacing": 2,
        "color_primary": "#1a1a1a",
        "color_secondary": "#555555",
        "color_accent": "#333333",
        "color_name": "#000000",
        "margins": {"top": 0.5, "bottom": 0.5, "left": 0.7, "right": 0.7},
        "section_separator": "line",
        "bullet_style": "–",
        "header_style": "centered",
    },
    "professional": {
        "name": "Professional",
        "description": "Traditional professional look with navy blue accents.",
        "font": "Helvetica",
        "heading_font": "Helvetica-Bold",
        "font_size": 10,
        "heading_size": 11,
        "name_size": 20,
        "section_spacing": 10,
        "bullet_spacing": 1.5,
        "color_primary": "#1a1a1a",
        "color_secondary": "#555555",
        "color_accent": "#1B3A5C",
        "color_name": "#1B3A5C",
        "margins": {"top": 0.5, "bottom": 0.5, "left": 0.65, "right": 0.65},
        "section_separator": "line",
        "bullet_style": "•",
        "header_style": "centered",
    },
    "modern": {
        "name": "Modern",
        "description": "Contemporary design with teal accents and clean typography.",
        "font": "Helvetica",
        "heading_font": "Helvetica-Bold",
        "font_size": 10,
        "heading_size": 11,
        "name_size": 22,
        "section_spacing": 10,
        "bullet_spacing": 1.5,
        "color_primary": "#2C3E50",
        "color_secondary": "#7F8C8D",
        "color_accent": "#0E8A6E",
        "color_name": "#0E8A6E",
        "margins": {"top": 0.45, "bottom": 0.45, "left": 0.65, "right": 0.65},
        "section_separator": "line",
        "bullet_style": "▸",
        "header_style": "centered",
    },
}


def get_template_config(template_name: TemplateName) -> dict[str, Any]:
    """Get configuration for a named template.

    Args:
        template_name: One of 'minimal', 'professional', 'modern'.

    Returns:
        Template configuration dict.
    """
    if template_name not in TEMPLATE_CONFIGS:
        raise ValueError(
            f"Unknown template '{template_name}'. "
            f"Available: {list(TEMPLATE_CONFIGS.keys())}"
        )
    return TEMPLATE_CONFIGS[template_name]


def list_templates() -> list[dict[str, str]]:
    """List all available templates with descriptions.

    Returns:
        List of dicts with 'name' and 'description'.
    """
    return [
        {"name": name, "description": conf["description"]}
        for name, conf in TEMPLATE_CONFIGS.items()
    ]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Templated PDF Export
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def export_pdf_template(
    resume_data: dict[str, Any],
    template_name: TemplateName = "professional",
    output_path: str | None = None,
) -> Path:
    """Export resume to PDF using a specific template.

    Args:
        resume_data: Structured resume dictionary.
        template_name: Template to use.
        output_path: Optional custom output path.

    Returns:
        Path to the generated PDF.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether,
    )
    from utils.file_export import generate_filename

    config = get_template_config(template_name)
    margins = config["margins"]

    if output_path is None:
        fname = generate_filename(resume_data.get("name", "resume"), "pdf")
        output_path = str(OUTPUTS_DIR / fname)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        topMargin=margins["top"] * inch,
        bottomMargin=margins["bottom"] * inch,
        leftMargin=margins["left"] * inch,
        rightMargin=margins["right"] * inch,
    )

    styles = getSampleStyleSheet()
    fs = config["font_size"]
    sec_space = config.get("section_spacing", 10)

    # ── Custom styles based on template ──
    styles.add(ParagraphStyle(
        name="TplName",
        parent=styles["Title"],
        fontSize=config["name_size"],
        leading=config["name_size"] + 6,
        textColor=HexColor(config.get("color_name", config["color_accent"])),
        alignment=TA_CENTER,
        spaceAfter=2,
        spaceBefore=0,
        fontName=config["heading_font"],
    ))
    styles.add(ParagraphStyle(
        name="TplContact",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        textColor=HexColor(config["color_secondary"]),
        alignment=TA_CENTER,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="TplHeading",
        parent=styles["Heading2"],
        fontSize=config["heading_size"],
        leading=config["heading_size"] + 3,
        textColor=HexColor(config["color_accent"]),
        spaceBefore=sec_space,
        spaceAfter=3,
        fontName=config["heading_font"],
        keepWithNext=True,
    ))
    styles.add(ParagraphStyle(
        name="TplBody",
        parent=styles["Normal"],
        fontSize=fs,
        leading=fs + 3.5,
        textColor=HexColor(config["color_primary"]),
        spaceAfter=2,
        fontName=config["font"],
        alignment=TA_JUSTIFY,
    ))
    styles.add(ParagraphStyle(
        name="TplJobTitle",
        parent=styles["Normal"],
        fontSize=fs + 0.5,
        leading=fs + 4,
        textColor=HexColor(config["color_primary"]),
        spaceAfter=0,
        spaceBefore=6,
        fontName=config["heading_font"],
    ))
    styles.add(ParagraphStyle(
        name="TplMeta",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        textColor=HexColor(config["color_secondary"]),
        spaceAfter=2,
        fontName=config["font"],
    ))
    styles.add(ParagraphStyle(
        name="TplBullet",
        parent=styles["Normal"],
        fontSize=fs,
        leading=fs + 3,
        textColor=HexColor(config["color_primary"]),
        leftIndent=16,
        spaceAfter=config.get("bullet_spacing", 1.5),
        bulletIndent=4,
        fontName=config["font"],
    ))
    styles.add(ParagraphStyle(
        name="TplSkills",
        parent=styles["Normal"],
        fontSize=fs,
        leading=fs + 4,
        textColor=HexColor(config["color_primary"]),
        spaceAfter=2,
        fontName=config["font"],
    ))

    elements: list[Any] = []
    bullet = config["bullet_style"]

    def add_divider():
        if config["section_separator"] == "line":
            elements.append(
                HRFlowable(
                    width="100%",
                    thickness=0.5,
                    color=HexColor(config["color_accent"]),
                    spaceAfter=4,
                    spaceBefore=0,
                )
            )
        elif config["section_separator"] == "space":
            elements.append(Spacer(1, 6))

    def _esc(text: str) -> str:
        """Escape XML special characters for ReportLab Paragraph."""
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

    # ── Header ──
    elements.append(Paragraph(_esc(resume_data.get("name", "Your Name")), styles["TplName"]))
    contact_parts = [
        resume_data[k] for k in ("email", "phone", "location", "linkedin", "website")
        if resume_data.get(k)
    ]
    if contact_parts:
        elements.append(Paragraph("  ·  ".join(_esc(p) for p in contact_parts), styles["TplContact"]))

    # ── Summary ──
    if resume_data.get("summary"):
        elements.append(Paragraph("PROFESSIONAL SUMMARY", styles["TplHeading"]))
        add_divider()
        elements.append(Paragraph(_esc(resume_data["summary"]), styles["TplBody"]))

    # ── Skills ──
    if resume_data.get("skills"):
        elements.append(Paragraph("SKILLS", styles["TplHeading"]))
        add_divider()
        skills = resume_data["skills"]
        # Group into rows of ~8 skills each for better readability
        chunk_size = 8
        for i in range(0, len(skills), chunk_size):
            chunk = skills[i:i+chunk_size]
            skills_text = f"  {bullet}  ".join(_esc(s) for s in chunk)
            elements.append(Paragraph(skills_text, styles["TplSkills"]))

    # ── Experience ──
    if resume_data.get("experience"):
        elements.append(Paragraph("EXPERIENCE", styles["TplHeading"]))
        add_divider()
        for exp in resume_data["experience"]:
            exp_elements: list[Any] = []
            title = exp.get('title', '')
            company = exp.get('company', '')
            title_text = f"<b>{_esc(title)}</b>"
            if company:
                title_text += f"  –  {_esc(company)}"
            exp_elements.append(Paragraph(title_text, styles["TplJobTitle"]))

            date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            loc = exp.get("location", "")
            meta = f"{_esc(loc)}  |  {_esc(date_str)}" if loc else _esc(date_str)
            exp_elements.append(Paragraph(meta, styles["TplMeta"]))

            for b in exp.get("bullets", []):
                exp_elements.append(Paragraph(f"{bullet}  {_esc(b)}", styles["TplBullet"]))

            # Keep job title + first few bullets together on same page
            elements.append(KeepTogether(exp_elements[:4]))
            elements.extend(exp_elements[4:])

    # ── Education ──
    if resume_data.get("education"):
        elements.append(Paragraph("EDUCATION", styles["TplHeading"]))
        add_divider()
        for edu in resume_data["education"]:
            deg_text = f"<b>{_esc(edu.get('degree', ''))}</b>"
            inst = edu.get('institution', '')
            if inst:
                deg_text += f"  –  {_esc(inst)}"
            elements.append(Paragraph(deg_text, styles["TplJobTitle"]))
            meta_parts = []
            if edu.get("location"):
                meta_parts.append(_esc(edu["location"]))
            if edu.get("graduation_date"):
                meta_parts.append(_esc(edu["graduation_date"]))
            if edu.get("gpa"):
                meta_parts.append(f"GPA: {_esc(edu['gpa'])}")
            if meta_parts:
                elements.append(Paragraph("  |  ".join(meta_parts), styles["TplMeta"]))

    # ── Certifications ──
    if resume_data.get("certifications"):
        elements.append(Paragraph("CERTIFICATIONS", styles["TplHeading"]))
        add_divider()
        for cert in resume_data["certifications"]:
            elements.append(Paragraph(f"{bullet}  {_esc(cert)}", styles["TplBullet"]))

    # ── Projects ──
    if resume_data.get("projects"):
        elements.append(Paragraph("PROJECTS", styles["TplHeading"]))
        add_divider()
        for proj in resume_data["projects"]:
            proj_name = _esc(proj.get('name', ''))
            proj_tech = proj.get('technologies', '').strip()
            txt = f"<b>{proj_name}</b>"
            if proj_tech:
                txt += f"  <i>({_esc(proj_tech)})</i>"
            elements.append(Paragraph(txt, styles["TplJobTitle"]))
            if proj.get("description"):
                elements.append(Paragraph(_esc(proj["description"]), styles["TplBody"]))

    # ── Languages ──
    if resume_data.get("languages"):
        elements.append(Paragraph("LANGUAGES", styles["TplHeading"]))
        add_divider()
        elements.append(Paragraph(
            f"  {bullet}  ".join(_esc(l) for l in resume_data["languages"]),
            styles["TplSkills"],
        ))

    # ── Research Papers ──
    if resume_data.get("research_papers"):
        elements.append(Paragraph("RESEARCH & PUBLICATIONS", styles["TplHeading"]))
        add_divider()
        for paper in resume_data["research_papers"]:
            title = _esc(paper.get('title', ''))
            authors = _esc(paper.get('authors', ''))
            venue = _esc(paper.get('venue', ''))
            year = _esc(paper.get('year', ''))
            url = paper.get('url', '').strip()
            txt = f"<b>{title}</b>"
            if authors:
                txt += f" — {authors}"
            meta = []
            if venue:
                meta.append(f"<i>{venue}</i>")
            if year:
                meta.append(year)
            if url:
                meta.append(f'<a href="{url}" color="#0969DA">{_esc(url[:50])}</a>')
            if meta:
                txt += f"<br/>{' | '.join(meta)}"
            elements.append(Paragraph(txt, styles["TplBody"]))

    # ── Volunteer Work ──
    if resume_data.get("volunteer_work"):
        elements.append(Paragraph("VOLUNTEER EXPERIENCE", styles["TplHeading"]))
        add_divider()
        for vol in resume_data["volunteer_work"]:
            role = _esc(vol.get('role', ''))
            org = _esc(vol.get('organization', ''))
            txt = f"<b>{role}</b>"
            if org:
                txt += f" — {org}"
            elements.append(Paragraph(txt, styles["TplJobTitle"]))
            dates = []
            if vol.get('start_date'):
                dates.append(vol['start_date'])
            if vol.get('end_date'):
                dates.append(vol['end_date'])
            if dates:
                elements.append(Paragraph(' — '.join(dates), styles["TplMeta"]))
            if vol.get('description'):
                elements.append(Paragraph(_esc(vol['description']), styles["TplBody"]))

    # ── Achievements ──
    if resume_data.get("achievements"):
        elements.append(Paragraph("ACHIEVEMENTS & AWARDS", styles["TplHeading"]))
        add_divider()
        for ach in resume_data["achievements"]:
            elements.append(Paragraph(f"{bullet}  {_esc(ach)}", styles["TplBullet"]))

    doc.build(elements)
    logger.info("Template PDF exported (%s): %s", template_name, output_path)
    return Path(output_path)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Templated DOCX Export
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def export_docx_template(
    resume_data: dict[str, Any],
    template_name: TemplateName = "professional",
    output_path: str | None = None,
) -> Path:
    """Export resume to DOCX using a specific template.

    Args:
        resume_data: Structured resume dictionary.
        template_name: Template to use.
        output_path: Optional custom output path.

    Returns:
        Path to the generated DOCX.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree
    from utils.file_export import generate_filename

    config = get_template_config(template_name)
    margins = config["margins"]

    if output_path is None:
        fname = generate_filename(resume_data.get("name", "resume"), "docx")
        output_path = str(OUTPUTS_DIR / fname)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(margins["top"])
        section.bottom_margin = Inches(margins["bottom"])
        section.left_margin = Inches(margins["left"])
        section.right_margin = Inches(margins["right"])

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(config["font_size"])
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.space_before = Pt(0)

    def hex_to_rgb(hex_color: str) -> RGBColor:
        h = hex_color.lstrip("#")
        return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    primary = hex_to_rgb(config["color_primary"])
    secondary = hex_to_rgb(config["color_secondary"])
    accent = hex_to_rgb(config["color_accent"])
    name_color = hex_to_rgb(config.get("color_name", config["color_accent"]))
    bullet_char = config["bullet_style"]

    # ── Name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_para.paragraph_format.space_before = Pt(0)
    name_run = name_para.add_run(resume_data.get("name", "Your Name"))
    name_run.bold = True
    name_run.font.size = Pt(config["name_size"])
    name_run.font.color.rgb = name_color

    # ── Contact ──
    contact_parts = [
        resume_data[k] for k in ("email", "phone", "location", "linkedin", "website")
        if resume_data.get(k)
    ]
    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        cp.paragraph_format.space_before = Pt(0)
        cr = cp.add_run("  ·  ".join(contact_parts))
        cr.font.size = Pt(9)
        cr.font.color.rgb = secondary

    def add_heading(title: str):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(config.get("section_spacing", 10))
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(config["heading_size"])
        run.font.color.rgb = accent
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), config["color_accent"].lstrip("#"))

    def add_job_header(title_text: str, subtitle: str = ""):
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(5)
        tp.paragraph_format.space_after = Pt(0)
        tr = tp.add_run(title_text)
        tr.bold = True
        tr.font.size = Pt(config["font_size"] + 0.5)
        tr.font.color.rgb = primary
        if subtitle:
            sub_run = tp.add_run(f"  –  {subtitle}")
            sub_run.font.size = Pt(config["font_size"] + 0.5)
            sub_run.font.color.rgb = primary

    def add_meta(text: str):
        mp = doc.add_paragraph()
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after = Pt(2)
        mr = mp.add_run(text)
        mr.font.size = Pt(9)
        mr.font.color.rgb = secondary
        mr.italic = True

    def add_bullet(text: str):
        bp = doc.add_paragraph(f"{bullet_char} {text}")
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.left_indent = Inches(0.2)

    # ── Summary ──
    if resume_data.get("summary"):
        add_heading("Professional Summary")
        p = doc.add_paragraph(resume_data["summary"])
        p.paragraph_format.space_after = Pt(2)

    # ── Skills ──
    if resume_data.get("skills"):
        add_heading("Skills")
        p = doc.add_paragraph(f"  {bullet_char}  ".join(resume_data["skills"]))
        p.paragraph_format.space_after = Pt(2)

    # ── Experience ──
    if resume_data.get("experience"):
        add_heading("Experience")
        for exp in resume_data["experience"]:
            add_job_header(exp.get("title", ""), exp.get("company", ""))
            date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            loc = exp.get("location", "")
            meta = f"{loc}  |  {date_str}" if loc else date_str
            add_meta(meta)
            for b in exp.get("bullets", []):
                add_bullet(b)

    # ── Education ──
    if resume_data.get("education"):
        add_heading("Education")
        for edu in resume_data["education"]:
            add_job_header(edu.get("degree", ""), edu.get("institution", ""))
            parts = []
            if edu.get("location"):
                parts.append(edu["location"])
            if edu.get("graduation_date"):
                parts.append(edu["graduation_date"])
            if edu.get("gpa"):
                parts.append(f"GPA: {edu['gpa']}")
            if parts:
                add_meta("  |  ".join(parts))

    # ── Certifications ──
    if resume_data.get("certifications"):
        add_heading("Certifications")
        for cert in resume_data["certifications"]:
            add_bullet(cert)

    # ── Projects ──
    if resume_data.get("projects"):
        add_heading("Projects")
        for proj in resume_data["projects"]:
            proj_name = proj.get("name", "")
            proj_tech = proj.get("technologies", "").strip()
            add_job_header(proj_name)
            if proj_tech:
                tech_p = doc.add_paragraph()
                tech_p.paragraph_format.space_before = Pt(0)
                tech_p.paragraph_format.space_after = Pt(1)
                tech_r = tech_p.add_run(f"Technologies: {proj_tech}")
                tech_r.italic = True
                tech_r.font.size = Pt(9)
                tech_r.font.color.rgb = secondary
            if proj.get("description"):
                p = doc.add_paragraph(proj["description"])
                p.paragraph_format.space_after = Pt(2)

    # ── Languages ──
    if resume_data.get("languages"):
        add_heading("Languages")
        p = doc.add_paragraph(f"  {bullet_char}  ".join(resume_data["languages"]))
        p.paragraph_format.space_after = Pt(2)

    # ── Research Papers ──
    if resume_data.get("research_papers"):
        add_heading("Research & Publications")
        for paper in resume_data["research_papers"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run(paper.get('title', ''))
            r.bold = True
            r.font.size = Pt(10)
            authors = paper.get('authors', '').strip()
            if authors:
                a = p.add_run(f" — {authors}")
                a.font.size = Pt(9)
            meta = []
            if paper.get('venue'):
                meta.append(paper['venue'])
            if paper.get('year'):
                meta.append(paper['year'])
            if paper.get('url'):
                meta.append(paper['url'])
            if meta:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after = Pt(2)
                mr = mp.add_run(' | '.join(meta))
                mr.italic = True
                mr.font.size = Pt(9)
                mr.font.color.rgb = secondary

    # ── Volunteer Work ──
    if resume_data.get("volunteer_work"):
        add_heading("Volunteer Experience")
        for vol in resume_data["volunteer_work"]:
            add_job_header(vol.get('role', ''))
            meta = []
            if vol.get('organization'):
                meta.append(vol['organization'])
            dates = []
            if vol.get('start_date'):
                dates.append(vol['start_date'])
            if vol.get('end_date'):
                dates.append(vol['end_date'])
            if dates:
                meta.append(' — '.join(dates))
            if meta:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after = Pt(1)
                mr = mp.add_run(' | '.join(meta))
                mr.italic = True
                mr.font.size = Pt(9)
                mr.font.color.rgb = secondary
            if vol.get('description'):
                p = doc.add_paragraph(vol['description'])
                p.paragraph_format.space_after = Pt(2)

    # ── Achievements ──
    if resume_data.get("achievements"):
        add_heading("Achievements & Awards")
        for ach in resume_data["achievements"]:
            add_bullet(ach)

    doc.save(output_path)
    logger.info("Template DOCX exported (%s): %s", template_name, output_path)
    return Path(output_path)
